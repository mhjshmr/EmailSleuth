from flask import Flask, render_template, request, redirect, url_for, flash, send_file
from io import BytesIO
from urllib.parse import unquote
from docx import Document
from docx.shared import Pt
from Email_Analyzer import (
    load_email_from_bytes,
    extract_headers,
    parse_sender,
    analyze_sender,
    analyze_authentication,
    get_email_body_text,
    find_obfuscated_urls,
    extract_urls,
    analyze_url,
    extract_images,
    analyze_images,
    analyze_attachments,
    compute_threat_score,
    categorize_threat_score,
    generate_report_text,
    generate_risk_factors,
    derive_threat_types,
)


app = Flask(__name__)
app.secret_key = "change-this-to-a-random-secret"

@app.route("/", methods=["GET"])
def index():
    return render_template("index.html")

@app.route("/analyze", methods=["POST"])
def analyze():
    uploaded_file = request.files.get("eml_file")
    if not uploaded_file or uploaded_file.filename == "":
        flash("Please upload a .eml file.", "error")
        return redirect(url_for("index"))

    try:
        raw_data = uploaded_file.read()
        msg = load_email_from_bytes(raw_data)
    except Exception as exc:
        flash(f"Unable to parse email file: {exc}", "error")
        return redirect(url_for("index"))

    headers = extract_headers(msg)
    display_name, email_addr, domain = parse_sender(headers.get("From"))
    sender_info = (display_name, email_addr, domain)
    sender_findings = analyze_sender(display_name, email_addr, domain)
    auth_findings = analyze_authentication(headers)

    body_text, raw_html = get_email_body_text(msg)
    obf_candidates = find_obfuscated_urls(body_text)
    urls = extract_urls(body_text)
    for candidate in obf_candidates:
        if candidate not in urls:
            urls.append(candidate)

    url_analysis = {}
    for url in urls:
        hostname, findings = analyze_url(url, body_text)
        url_analysis[url] = (hostname, findings)

    images = extract_images(msg)
    image_findings, image_suspicious_count = analyze_images(images, raw_html)
    attachments_list, attachment_risk_count, attachment_findings = analyze_attachments(msg)
    attachment_score = min(attachment_risk_count * 10, 10)

    # Get score breakdown from compute_threat_score
    threat_score, auth_score, url_score, sender_score, image_score, attachment_score = compute_threat_score(
        headers,
        sender_info,
        auth_findings,
        url_analysis,
        image_suspicious_count,
        sender_findings,
        attachment_score,
    )
    if threat_score > 100:
        threat_score = 100

    risk_level = categorize_threat_score(threat_score)

    # Generate risk factors and threat types from the actual scores
    risk_factors = generate_risk_factors(
        sender_findings, auth_findings, url_analysis, image_findings, attachment_findings
    )
    threat_types = derive_threat_types(auth_score, url_score, sender_score, image_score)

    report_text = generate_report_text(
        headers,
        sender_info,
        sender_findings,
        auth_findings,
        urls,
        url_analysis,
        image_findings,
        attachments_list,
        attachment_findings,
        threat_score,
        risk_level,
    )

    return render_template(
        "result.html",
        headers=headers,
        sender_info=sender_info,
        sender_findings=sender_findings,
        auth_findings=auth_findings,
        urls=urls,
        url_analysis=url_analysis,
        image_findings=image_findings,
        image_suspicious_count=image_suspicious_count,
        attachments_list=attachments_list,
        attachment_findings=attachment_findings,
        threat_score=threat_score,
        risk_level=risk_level,
        auth_score=auth_score,
        url_score=url_score,
        sender_score=sender_score,
        image_score=image_score,
        risk_factors=risk_factors,
        threat_types=threat_types,
        report_text=report_text,
    )


@app.route("/download-report", methods=["POST"])
def download_report():
    """
    Generate a .docx report from the analysis results passed via form data.
    """
    # Get form data from the submitted form
    threat_score = request.form.get("threat_score", "N/A")
    risk_level = request.form.get("risk_level", "N/A")
    # Decode URL-encoded report text (newlines become %0A, etc.)
    report_text = unquote(request.form.get("report_text", "No report available."))
    filename = request.form.get("filename", "email_report")
    
    # Clean up filename - ensure it ends with .docx
    if not filename.endswith(".docx"):
        filename += ".docx"

    # Create a new Word document
    doc = Document()
    
    # Add title
    doc.add_heading("EmailSleuth - Phishing Analysis Report", 0)
    
    # Add threat score section
    doc.add_heading("Threat Assessment", level=1)
    doc.add_paragraph(f"Threat Score: {threat_score}/100")
    doc.add_paragraph(f"Risk Level: {risk_level}")
    
    # Add the full report text
    doc.add_heading("Detailed Report", level=1)
    # Split by lines and add each as a paragraph
    for line in report_text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)
    
    # Add footer
    doc.add_paragraph("\n" + "="*50)
    doc.add_paragraph("Generated by EmailSleuth - Email Header Forensics & Phishing Analyzer")

    # Save to in-memory buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Send as downloadable file
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)


